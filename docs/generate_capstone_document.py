from __future__ import annotations

from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "SOG_Capstone_System_Documentation.docx"
CONTEXT_IMG = DOCS / "capstone-context-flow-diagram.png"
DFD_IMG = DOCS / "capstone-dfd-level-1.png"

BLUE = "#608DB9"
DARK_BLUE = "#1F4D78"
INK = "#15233A"
MUTED = "#64748B"
LIGHT = "#EEF5FB"
BORDER = "#B8C9DA"
WHITE = "#FFFFFF"
GOLD = "#F5B811"


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Calibri Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Calibri.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = font(32, True)
FONT_H = font(23, True)
FONT_B = font(18)
FONT_BOLD = font(18, True)
FONT_SMALL = font(15)
FONT_SMALL_BOLD = font(15, True)


def wrap_text(text: str, fnt: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if fnt.getbbox(candidate)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines or [""]


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fnt, fill=INK):
    x1, y1, x2, y2 = box
    lines = wrap_text(text, fnt, x2 - x1 - 32)
    line_height = fnt.size + 6
    total = line_height * len(lines)
    y = y1 + ((y2 - y1) - total) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=fnt)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, font=fnt, fill=fill)
        y += line_height


def box(draw: ImageDraw.ImageDraw, xy, text, fill=WHITE, outline=BLUE, width=3, radius=22, fnt=FONT_BOLD):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)
    centered_text(draw, xy, text, fnt)


def rect(draw: ImageDraw.ImageDraw, xy, text, fill=WHITE, outline=BLUE, width=3, fnt=FONT_BOLD):
    draw.rectangle(xy, fill=fill, outline=outline, width=width)
    centered_text(draw, xy, text, fnt)


def store(draw: ImageDraw.ImageDraw, xy, text):
    x1, y1, x2, y2 = xy
    draw.rectangle(xy, fill="#F8FBFE", outline=DARK_BLUE, width=3)
    draw.line((x1 + 18, y1, x1 + 18, y2), fill=DARK_BLUE, width=3)
    centered_text(draw, (x1 + 18, y1, x2, y2), text, FONT_SMALL_BOLD, DARK_BLUE)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str = "", color=DARK_BLUE):
    draw.line((start[0], start[1], end[0], end[1]), fill=color, width=3)
    dx, dy = end[0] - start[0], end[1] - start[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    left = (-uy, ux)
    size = 14
    p1 = (end[0], end[1])
    p2 = (end[0] - ux * size + left[0] * size * 0.6, end[1] - uy * size + left[1] * size * 0.6)
    p3 = (end[0] - ux * size - left[0] * size * 0.6, end[1] - uy * size - left[1] * size * 0.6)
    draw.polygon([p1, p2, p3], fill=color)
    if label:
        mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
        lines = wrap_text(label, FONT_SMALL, 210)
        tw = max(FONT_SMALL.getbbox(line)[2] for line in lines) + 18
        th = len(lines) * (FONT_SMALL.size + 3) + 10
        draw.rounded_rectangle((mx - tw // 2, my - th // 2, mx + tw // 2, my + th // 2), radius=10, fill=WHITE, outline=BORDER, width=1)
        y = my - th // 2 + 5
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=FONT_SMALL)
            draw.text((mx - (bbox[2] - bbox[0]) / 2, y), line, font=FONT_SMALL, fill=MUTED)
            y += FONT_SMALL.size + 3


def make_context_diagram():
    img = Image.new("RGB", (1800, 1150), WHITE)
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((35, 35, 1765, 1115), radius=36, fill="#F7FAFD", outline="#D9E4EF", width=3)
    d.text((70, 70), "Context Flow Diagram", font=FONT_TITLE, fill=DARK_BLUE)
    d.text((70, 116), "SOG Glass and Aluminum Appointment Booking and Management System", font=FONT_B, fill=MUTED)

    system = (650, 430, 1150, 720)
    box(d, system, "SOG Appointment Booking\nand Management System", fill=LIGHT, outline=DARK_BLUE, width=4, radius=30, fnt=FONT_H)

    entities = {
        "Client / Customer": (105, 230, 420, 365),
        "Admin / Owner": (1380, 225, 1685, 360),
        "Worker / Staff": (1380, 785, 1685, 920),
        "PayPal Sandbox / API": (105, 790, 420, 925),
        "Email / SMS Gateway": (720, 885, 1080, 1020),
        "Google Maps Services": (680, 215, 1120, 350),
        "AR Measurement App": (110, 510, 430, 645),
    }
    for name, xy in entities.items():
        box(d, xy, name, fill=WHITE, outline=BLUE, radius=20, fnt=FONT_BOLD)

    arrow(d, (420, 298), (650, 500), "Booking details, quotation requests, OTP login")
    arrow(d, (650, 575), (420, 298), "Appointment status, quote, payment, work updates")
    arrow(d, (1380, 295), (1150, 500), "Manage users, products, schedules, reports")
    arrow(d, (1150, 570), (1380, 295), "Dashboards, audit logs, sales/payment data")
    arrow(d, (1380, 850), (1150, 640), "Work progress, remarks, photos")
    arrow(d, (1150, 680), (1380, 850), "Assigned appointments and work jobs")
    arrow(d, (420, 575), (650, 600), "Measured dimensions and selected 3D model")
    arrow(d, (650, 650), (420, 860), "Payment order request")
    arrow(d, (420, 860), (650, 690), "Payment confirmation / failure")
    arrow(d, (900, 885), (900, 720), "Email/SMS/notification delivery")
    arrow(d, (900, 350), (900, 430), "Map, route, geocoding")

    d.text((70, 1045), "Legend: rounded rectangles represent external actors/services and the central system; arrows show major information exchange.", font=FONT_SMALL, fill=MUTED)
    img.save(CONTEXT_IMG)


def make_dfd_diagram():
    img = Image.new("RGB", (2200, 1450), WHITE)
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((35, 35, 2165, 1415), radius=36, fill="#F7FAFD", outline="#D9E4EF", width=3)
    d.text((70, 70), "Data Flow Diagram - Level 1", font=FONT_TITLE, fill=DARK_BLUE)
    d.text((70, 116), "Clean Level 1 view of external actors, core processes, and persistent data stores", font=FONT_B, fill=MUTED)

    # Column headings
    d.text((105, 178), "EXTERNAL ACTORS", font=FONT_SMALL_BOLD, fill=MUTED)
    d.text((470, 178), "CORE SYSTEM PROCESSES", font=FONT_SMALL_BOLD, fill=MUTED)
    d.text((1270, 178), "DATA STORES", font=FONT_SMALL_BOLD, fill=MUTED)
    d.text((1710, 178), "EXTERNAL SERVICES", font=FONT_SMALL_BOLD, fill=MUTED)

    rows = [
        {
            "actor": "Customer, Admin,\nWorker",
            "process": "1.0\nAuthenticate Users\nand Manage Access",
            "store": "D1 Users,\nRoles,\nPermissions",
            "service": "OTP / Login\nSession",
            "in": "login request",
            "out": "account and role data",
        },
        {
            "actor": "Admin",
            "process": "2.0\nManage Product Catalog,\nOptions, Images, 3D Models",
            "store": "D2 Products,\nOptions,\nImages, 3D Models",
            "service": "3D Model\nStorage",
            "in": "product setup",
            "out": "catalog data",
        },
        {
            "actor": "Customer,\nAdmin",
            "process": "3.0\nBook and Manage\nAppointments",
            "store": "D3 Appointments,\nLocations,\nRemarks",
            "service": "Google Maps\nGeocoding",
            "in": "booking details",
            "out": "appointment record",
        },
        {
            "actor": "Admin, Worker,\nCustomer",
            "process": "4.0\nCreate, Review,\nApprove and Sign\nQuotations",
            "store": "D4 Quotations,\nItems, Photos,\nSignatures",
            "service": "PDF and\nE-signature",
            "in": "quote actions",
            "out": "quotation data",
        },
        {
            "actor": "Admin,\nWorker",
            "process": "5.0\nSchedule Work Jobs\nand Back Jobs",
            "store": "D5 Work Jobs,\nBack Jobs,\nRemarks",
            "service": "Calendar\nSchedule",
            "in": "assignment updates",
            "out": "job records",
        },
        {
            "actor": "Customer,\nAdmin",
            "process": "6.0\nProcess Payments,\nCharges and Sales",
            "store": "D6 Payments,\nCharges,\nSales Reports",
            "service": "PayPal\nSandbox / API",
            "in": "payment request",
            "out": "payment status",
        },
        {
            "actor": "All Roles",
            "process": "7.0\nNotify, Audit and\nGenerate Reports",
            "store": "D7 Notifications,\nAudit Logs,\nReports",
            "service": "Email / SMS\nGateway",
            "in": "system events",
            "out": "notifications / reports",
        },
    ]

    actor_x = (90, 225, 345, 325)
    process_x = (455, 225, 925, 325)
    store_x = (1185, 225, 1485, 325)
    service_x = (1735, 225, 2045, 325)
    row_gap = 163

    process_centers: list[tuple[int, int]] = []
    for i, row in enumerate(rows):
        y_offset = i * row_gap
        actor_xy = (actor_x[0], actor_x[1] + y_offset, actor_x[2], actor_x[3] + y_offset)
        proc_xy = (process_x[0], process_x[1] + y_offset, process_x[2], process_x[3] + y_offset)
        store_xy = (store_x[0], store_x[1] + y_offset, store_x[2], store_x[3] + y_offset)
        service_xy = (service_x[0], service_x[1] + y_offset, service_x[2], service_x[3] + y_offset)

        box(d, actor_xy, row["actor"], fill=WHITE, outline=BLUE, radius=18, fnt=FONT_SMALL_BOLD)
        rect(d, proc_xy, row["process"], fill=LIGHT, outline=DARK_BLUE, fnt=FONT_SMALL_BOLD)
        store(d, store_xy, row["store"])
        box(d, service_xy, row["service"], fill=WHITE, outline=BLUE, radius=18, fnt=FONT_SMALL_BOLD)

        cy = (proc_xy[1] + proc_xy[3]) // 2
        process_centers.append(((proc_xy[0] + proc_xy[2]) // 2, cy))
        arrow(d, (actor_xy[2], cy), (proc_xy[0], cy), row["in"])
        arrow(d, (proc_xy[2], cy), (store_xy[0], cy), row["out"])
        arrow(d, (store_xy[2], cy), (service_xy[0], cy), "service data")
        arrow(d, (service_xy[0], cy + 26), (store_xy[2], cy + 26), "response")

    # Vertical business flow through the processes.
    for start, end in zip(process_centers, process_centers[1:]):
        arrow(d, (start[0], start[1] + 50), (end[0], end[1] - 50), "next workflow event")

    d.text((70, 1372), "DFD note: each row shows who sends data, the system process that handles it, where records are stored, and the outside service involved when applicable.", font=FONT_SMALL, fill=MUTED)
    img.save(DFD_IMG)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "D9E4EF")
        borders.append(tag)
    tbl_pr.append(borders)


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.add_run(text)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    set_table_borders(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        shade_cell(cell, "#F2F4F7")
        cell.width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()
    return table


def configure_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(21, 35, 58)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, DARK_BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_doc():
    make_context_diagram()
    make_dfd_diagram()

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOG Glass and Aluminum Services\nAppointment Booking and Management System")
    run.font.name = "Calibri"
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE.replace("#", ""))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Capstone Technical Background, Resources, Scope, Context Flow Diagram, and Data Flow Diagram")
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Prepared for: College Capstone Documentation")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)

    add_heading(doc, "1. Technical Background", 1)
    doc.add_paragraph(
        "SOG Glass and Aluminum Services requires a reliable digital system for handling customer appointment requests, product quotations, work scheduling, payment records, and job progress monitoring. In a traditional manual process, booking requests, quotation details, customer contacts, measurements, payment follow-ups, and worker assignments can become scattered across messages, paper notes, and informal reminders. This creates risks such as missed appointments, inconsistent quotation records, delayed updates, and limited visibility for customers and staff."
    )
    doc.add_paragraph(
        "The proposed system is a web-based appointment booking and management platform designed to centralize the company workflow. It supports online appointment booking, product quotation preparation, customer tracking, staff scheduling, work job monitoring, payment recording, audit logs, and real-time notifications. It also includes a separate augmented reality measurement tool that uses WebXR and Three.js to help customers or staff visualize and capture approximate product dimensions using supported mobile devices."
    )
    doc.add_paragraph(
        "The system follows a client-server architecture. The frontend provides user interfaces for customers, administrators, and workers. The backend exposes secure API endpoints, handles validation and business rules, stores system records, and triggers notifications. The database stores users, products, appointments, quotations, work jobs, payments, audits, and related records."
    )

    add_heading(doc, "2. Technologies to be Used", 1)
    add_heading(doc, "2.1 Frontend Technologies", 2)
    add_table(
        doc,
        ["Technology", "Use in the System"],
        [
            ["Next.js 16", "Main customer, admin, and worker web application with routed pages and server/client rendering support."],
            ["React 19", "Reusable component-based user interface for dashboards, forms, modals, quotations, and customer pages."],
            ["TypeScript", "Static typing for frontend models, API payloads, form data, and reusable components."],
            ["Tailwind CSS 4", "Utility-based styling system for responsive and consistent layouts."],
            ["shadcn UI / Radix UI", "Accessible UI primitives such as dialogs, drawers, select controls, buttons, tabs, and form components."],
            ["Lucide React", "Consistent icon library for navigation, actions, status indicators, and dashboard cards."],
            ["Zod", "Client-side schema validation for forms and safer typed data handling."],
            ["FullCalendar", "Appointment and worker schedule calendar views."],
            ["Recharts", "Dashboard and sales chart visualizations."],
            ["Google Maps", "Customer location display, route reference, and address-based map views."],
            ["@google/model-viewer", "Admin preview of uploaded GLB product 3D models."],
            ["PayPal React SDK", "Customer payment buttons for PayPal sandbox and card-supported checkout flows."],
            ["Laravel Echo and Pusher JS", "Realtime notification and page refresh events from Laravel Reverb broadcasting."],
            ["Playwright", "Frontend end-to-end testing for major user flows."],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "2.2 Backend Technologies", 2)
    add_table(
        doc,
        ["Technology", "Use in the System"],
        [
            ["Laravel 13", "Main backend API framework for routing, validation, services, events, listeners, notifications, and data resources."],
            ["PHP 8.3", "Backend programming language used by Laravel."],
            ["Laravel Sanctum", "API authentication for logged-in customer and staff requests."],
            ["Laravel Fortify", "Staff profile, password, and two-factor authentication features."],
            ["Spatie Laravel Permission", "Role and permission management for admin, sub-admin, worker, and customer access."],
            ["Laravel Reverb", "Self-hosted realtime broadcasting for notifications and live status updates."],
            ["Laravel Notifications", "Database notifications for appointment, quotation, work job, payment, and status changes."],
            ["Laravel Auditing", "Audit trail for important create, update, and delete actions."],
            ["Maatwebsite Laravel Excel", "Excel and CSV exports for reports such as sales records."],
            ["Spatie Laravel PDF / Browsershot", "Quotation and report PDF generation."],
            ["PayPal REST API", "Payment order creation, capture, and payment record synchronization."],
            ["Relational Database", "Persistent storage for users, products, appointments, quotations, work jobs, payments, notifications, audits, and supporting records."],
        ],
        [1.9, 4.6],
    )

    add_heading(doc, "2.3 Augmented Reality and 3D Technologies", 2)
    add_table(
        doc,
        ["Technology", "Use in the System"],
        [
            ["Vite + React", "Separate lightweight AR frontend application for faster mobile testing and deployment."],
            ["Three.js", "3D rendering engine for GLB models, anchor points, measurement lines, labels, and object placement."],
            ["WebXR Hit Testing", "Detects real-world surfaces and places points or models in AR space on supported Android/Chrome devices."],
            ["GLB 3D Models", "Product models uploaded and linked to catalog items for AR visualization."],
            ["Blender", "Creation and preparation of 3D product models before upload."],
        ],
        [1.9, 4.6],
    )

    add_heading(doc, "3. System Resources", 1)
    add_heading(doc, "3.1 Hardware Resources", 2)
    add_table(
        doc,
        ["Hardware", "Specification / Purpose"],
        [
            ["Development Laptop", "MacBook Air M1 with 24 GB unified memory and 512 GB storage."],
            ["Mobile Test Device", "AR-capable Android smartphone with Chrome/WebXR support for AR hit testing."],
            ["Internet Connection", "Required for API access, PayPal sandbox testing, maps, notifications, and deployment testing."],
            ["Optional External Display", "Helpful for development, debugging, documentation, and browser testing."],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "3.2 Software Resources", 2)
    add_table(
        doc,
        ["Software", "Purpose"],
        [
            ["Visual Studio Code", "Primary code editor for frontend, backend, AR, and documentation work."],
            ["Laravel Herd", "Local PHP/Laravel development environment."],
            ["Node.js and npm", "Frontend, AR application, and testing dependency management."],
            ["Composer", "Laravel package and dependency management."],
            ["Google Chrome / Chrome DevTools", "Browser testing, WebXR testing, network inspection, and UI debugging."],
            ["Blender", "Creation, optimization, and export of product 3D models in GLB format."],
            ["Microsoft Word", "Capstone document preparation and formatting."],
            ["Git", "Version control for source code and documentation changes."],
            ["PayPal Developer Sandbox", "Safe testing of PayPal checkout without using real money."],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "4. Scope of the Study", 1)
    doc.add_paragraph("The study covers the design and development of a web-based appointment booking and business management system for SOG Glass and Aluminum Services. The system scope includes:")
    add_bullets(
        doc,
        [
            "Customer appointment booking with or without quotation items.",
            "Customer OTP login using email or phone and customer dashboard access.",
            "Product catalog management, including product images, variants, option groups, options, categories, and uploaded 3D product models.",
            "Quotation creation, item status management, digital customer review, e-signature, approved-item PDF generation, and re-signing when approved items are changed.",
            "Admin appointment management, including confirmation, rescheduling, cancellation, reopening, no-show handling, and calendar viewing.",
            "Worker assignment, worker dashboard, assigned appointment/work job views, status updates, remarks, and before/after photo visibility.",
            "Work job and back job management, including jobs created from appointments and warranty or unfinished-work scenarios.",
            "Payment lifecycle support, including PayPal payments, down payments, full payments, manual/cash records, additional charges, payment status tracking, and sales reports.",
            "Realtime notification updates using Laravel Reverb and database notifications.",
            "Audit logging for security and accountability.",
            "AR measurement and product model visualization using a separate React/Vite WebXR application.",
            "Dashboard analytics, sales reporting, calendar monitoring, PDF exports, Excel/CSV exports, and frontend E2E testing.",
        ],
    )

    add_heading(doc, "5. Limitations of the Study", 1)
    add_bullets(
        doc,
        [
            "The AR measurement tool provides approximate measurements and does not replace professional on-site measurement by company personnel.",
            "WebXR hit testing depends on supported devices, supported browsers, camera permission, lighting, surface texture, and browser security requirements such as HTTPS.",
            "The system requires an internet connection for online booking, notifications, maps, PayPal checkout, and real-time updates.",
            "PayPal integration is tested through sandbox mode; live payment processing requires correct live credentials and PayPal account approval.",
            "SMS and email delivery depend on configured third-party providers and network availability.",
            "The system does not cover payroll, supplier purchasing, full inventory accounting, or manufacturing automation.",
            "3D model accuracy depends on the quality and scale of uploaded GLB files prepared in Blender or other modeling software.",
            "Real-time updates require the Laravel Reverb server and frontend Echo connection to be running properly.",
            "Final quotation prices and job completion still require staff validation because actual site conditions may change after booking.",
        ],
    )

    add_heading(doc, "6. Context Flow Diagram", 1)
    doc.add_paragraph(
        "The context flow diagram shows the system as one central process and identifies how external actors and services exchange information with it."
    )
    doc.add_picture(str(CONTEXT_IMG), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph("Figure 1. Context Flow Diagram of the SOG Appointment Booking and Management System.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "7. Data Flow Diagram", 1)
    doc.add_paragraph(
        "The Level 1 DFD decomposes the system into major processes and shows how data moves between customers, staff, external services, and internal data stores."
    )
    doc.add_picture(str(DFD_IMG), width=Inches(6.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = doc.add_paragraph("Figure 2. Level 1 Data Flow Diagram of the SOG System.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    add_heading(doc, "7.1 DFD Process Descriptions", 2)
    add_table(
        doc,
        ["Process", "Description"],
        [
            ["1.0 Authenticate Users and Manage Access", "Handles customer OTP login, staff login, two-factor authentication, roles, permissions, and session access."],
            ["2.0 Manage Product Catalog", "Stores product information, categories, variants, option groups, images, and 3D model files."],
            ["3.0 Book and Manage Appointments", "Receives booking requests, stores preferred or confirmed schedules, and supports status changes."],
            ["4.0 Create, Review, and Sign Quotations", "Creates quotation items, applies product/options snapshots, manages approval statuses, and stores digital signatures."],
            ["5.0 Schedule Work Jobs and Back Jobs", "Creates work jobs from appointments, assigns workers, tracks progress, and supports back jobs for warranty or unfinished work."],
            ["6.0 Process Payments, Charges, and Sales", "Creates PayPal payment records, tracks manual/cash payments, handles additional charges, and supports sales reports."],
            ["7.0 Notify, Audit, and Generate Reports", "Creates notifications, broadcasts real-time events, stores audit records, and produces operational reports."],
        ],
        [2.2, 4.3],
    )

    add_heading(doc, "8. Summary", 1)
    doc.add_paragraph(
        "The proposed system supports the full operational cycle of SOG Glass and Aluminum Services from customer inquiry to appointment booking, quotation preparation, worker scheduling, payment tracking, job completion, and reporting. By combining Laravel, Next.js, WebXR, PayPal, real-time notifications, audit logs, and role-based access control, the system provides a modern platform that improves transparency, traceability, and customer service while still allowing staff to validate final measurements and pricing."
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("SOG Glass and Aluminum Services Appointment Booking and Management System")
    rr.font.size = Pt(8.5)
    rr.font.color.rgb = RGBColor.from_string(MUTED.replace("#", ""))

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
